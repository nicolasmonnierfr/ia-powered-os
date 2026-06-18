#!/usr/bin/env python3
"""
desanonymiser.py — REPERSONNALISATION (chemin inverse de l'anonymisation, #12).

Les rapports d'analyse produits à partir d'un transcript anonymisé contiennent
les pseudos (PERSONNE_1, SOCIETE_1…). Pour les livrer au dirigeant, ce script
réinjecte les vrais noms à partir de la mémoire client (`memoire_client.json`).

Le retour est NON AMBIGU : chaque pseudo est unique -> un seul remplacement
possible (contrairement à l'aller, où plusieurs variantes -> un pseudo).

⚠️ SÉCURITÉ : le fichier produit CONTIENT À NOUVEAU LES VRAIS NOMS. Il ne doit
JAMAIS être renvoyé à une IA externe. La sortie est nommée «_REPERSONNALISE».

Cible du remplacement : par défaut le `canonique` (forme la plus complète).
Option --court pour viser la variante la plus courte.

Formats : .txt / .md / .srt traités nativement. .docx via le skill docx
(python-docx) si disponible.

100% local. Voir memoire.py (schéma) et BACKLOG.md #12.

Usage :
    python desanonymiser.py rapport.md --memoire memoire_client.json
    python desanonymiser.py rapport.txt --memoire m.json --court
    python desanonymiser.py rapport.docx --memoire m.json
"""

import argparse
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import memoire as M  # noqa: E402

# Console Windows en cp1252 : un print accentué/emoji (« ⚠ ») planterait
# (UnicodeEncodeError) et ferait sortir le script en erreur -> ia repersonnaliser
# le verrait comme un échec alors que le fichier est produit. On force UTF-8.
for _flux in (sys.stdout, sys.stderr):
    try:
        _flux.reconfigure(encoding="utf-8")
    except Exception:
        pass


def die(msg, code=1):
    print(f"[ERREUR] {msg}", file=sys.stderr)
    sys.exit(code)


def _est_placeholder(s):
    """Étiquette technique du tagueur, PAS un vrai nom : « NON_AFFECTE » ou
    « Locuteur N ». A ne jamais réinjecter telle quelle dans un livrable."""
    return (not s) or s == "NON_AFFECTE" or bool(re.match(r"(?i)\s*locuteur\s*\d+\s*$", s))


def construire_mapping(mem, court=False):
    """Liste (pseudo, remplacement), pseudos longs d'abord (PERSONNE_10 avant _1).
    On écarte les variantes/canoniques PLACEHOLDER (NON_AFFECTE, « Locuteur N ») :
    un locuteur jamais nommé ne doit pas réintroduire ces étiquettes techniques.
    En dernier recours (aucune vraie variante), on conserve le PSEUDO (signale un
    nom à compléter dans la mémoire, au lieu d'un faux « Locuteur 2 »)."""
    out = []
    for e in mem.get("entrees", []):
        vraies = [v for v in (e.get("variantes") or []) if not _est_placeholder(v)]
        cano = e.get("canonique")
        if court and vraies:
            repl = min(vraies, key=len)
        elif cano and not _est_placeholder(cano):
            repl = cano
        elif vraies:
            repl = max(vraies, key=len)
        else:
            repl = e["pseudo"]
        out.append((e["pseudo"], repl))
    out.sort(key=lambda t: -len(t[0]))
    return out


def make_replacer(mapping):
    compiled = []
    for pseudo, repl in mapping:
        # mot entier : le pseudo est alphanumérique + underscore
        pat = re.compile(r"\b" + re.escape(pseudo) + r"\b")
        compiled.append((pat, repl))
    counts = {}

    def replace_in(text):
        for pat, repl in compiled:
            def _sub(m, _r=repl, _p=pat.pattern):
                counts[_p] = counts.get(_p, 0) + 1
                return _r
            text = pat.sub(_sub, text)
        return text
    return replace_in, counts


# ---------------------------------------------------------------------------
# Traitement par format
# ---------------------------------------------------------------------------
def traiter_texte(path, replace_in):
    raw = path.read_text(encoding="utf-8")
    return replace_in(raw)


def traiter_docx(path, out_path, replace_in):
    """Remplace dans un .docx en préservant la mise en forme (run par run)."""
    try:
        from docx import Document
    except ImportError:
        die("python-docx requis pour les .docx (voir skill docx / bootstrap).")
    doc = Document(str(path))

    def process_paragraph(p):
        for run in p.runs:
            if run.text:
                run.text = replace_in(run.text)

    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
    doc.save(str(out_path))


def main():
    ap = argparse.ArgumentParser(
        description="Repersonnalise un rapport anonymisé (réinjecte les vrais noms).")
    ap.add_argument("rapport", help="Fichier .txt/.md/.srt/.docx contenant des pseudos.")
    ap.add_argument("--memoire", required=True, help=f"{M.NOM_MEMOIRE} (table pseudo<->réel).")
    ap.add_argument("--court", action="store_true",
                    help="Remplacer par la variante la plus courte (défaut : canonique).")
    ap.add_argument("--out", help="Fichier de sortie (défaut : <nom>_REPERSONNALISE<ext>).")
    args = ap.parse_args()

    rpath = Path(args.rapport)
    if not rpath.exists():
        die(f"Rapport introuvable : {rpath}")
    mpath = Path(args.memoire)
    if not mpath.exists():
        die(f"Mémoire introuvable : {mpath}")

    mem = M.charger_memoire(mpath)
    if not mem.get("entrees"):
        die("La mémoire ne contient aucune entrée (rien à repersonnaliser).")

    mapping = construire_mapping(mem, court=args.court)
    replace_in, counts = make_replacer(mapping)

    out = Path(args.out) if args.out else rpath.with_name(
        f"{rpath.stem}_REPERSONNALISE{rpath.suffix}")

    ext = rpath.suffix.lower()
    if ext == ".docx":
        traiter_docx(rpath, out, replace_in)
    elif ext in (".txt", ".md", ".srt", ""):
        out.write_text(traiter_texte(rpath, replace_in), encoding="utf-8")
    else:
        die(f"Format non pris en charge : {ext}. Formats : .txt .md .srt .docx")

    total = sum(counts.values())
    print(f"[OK] Fichier repersonnalisé : {out}")
    print(f"     {total} pseudo(s) remplacé(s).")
    # pseudos jamais rencontrés
    vus = {p.replace("\\b", "").replace("\\", "") for p in counts}  # approximatif
    print("\n⚠ Ce fichier contient les VRAIS NOMS : ne JAMAIS l'envoyer à une IA externe.")


if __name__ == "__main__":
    main()
